"""Unit tests for the DOCXProcessor module.

Cover paragraph/table extraction, metadata enrichment, header/footer inclusion,
and error handling for invalid or corrupted DOCX files.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable
from unittest.mock import MagicMock, patch
from zipfile import BadZipFile

import pytest
from docx.opc.exceptions import PackageNotFoundError

from zapomni_core.exceptions import ProcessingError, ValidationError
from zapomni_core.processors.docx_processor import DOCXProcessor


def _create_docx_path(tmp_path: Path, name: str = "document.docx") -> Path:
    path = tmp_path / name
    path.write_bytes(b"PK\x03\x04")
    return path


def _build_paragraphs(texts: Iterable[str]) -> list[MagicMock]:
    return [MagicMock(text=text) for text in texts]


def _build_table(rows: Iterable[Iterable[str]]) -> MagicMock:
    table = MagicMock()
    table.rows = []
    for row in rows:
        row_mock = MagicMock()
        row_mock.cells = []
        for cell_text in row:
            cell = MagicMock()
            cell.text = cell_text
            row_mock.cells.append(cell)
        table.rows.append(row_mock)
    return table


def _build_core_properties(
    title: str | None = None,
    author: str | None = None,
    last_modified_by: str | None = None,
    modified: datetime | None = None,
) -> MagicMock:
    core = MagicMock()
    core.title = title
    core.author = author
    core.last_modified_by = last_modified_by
    core.modified = modified
    return core


class TestDOCXProcessor:
    """Tests for paragraph/table extraction and metadata handling."""

    def setup_method(self) -> None:
        self.processor = DOCXProcessor()

    @patch("zapomni_core.processors.docx_processor.Document")
    def test_process_docx_extracts_paragraphs_and_tables(
        self, mock_document: MagicMock, tmp_path: Path
    ) -> None:
        """Should clean text from paragraphs and tables."""
        doc_path = _create_docx_path(tmp_path)

        document = MagicMock()
        document.paragraphs = _build_paragraphs([" Hello world  ", "", "Paragraph two"])
        document.tables = [_build_table([["Cell 1", "Cell   2"], ["Row 2, Col 1", ""]])]
        document.sections = []
        document.core_properties = _build_core_properties()

        mock_document.return_value = document

        output = self.processor.process_docx(str(doc_path))

        assert "Hello world" in output
        assert "Paragraph two" in output
        assert "Cell 1 | Cell 2" in output
        assert "Row 2, Col 1" in output
        mock_document.assert_called_once_with(str(doc_path))

    @patch("zapomni_core.processors.docx_processor.Document")
    def test_process_docx_with_metadata_returns_data(
        self, mock_document: MagicMock, tmp_path: Path
    ) -> None:
        """Should return cleaned text and metadata dictionary."""
        doc_path = _create_docx_path(tmp_path)

        core = _build_core_properties(
            title="Report",
            author=None,
            last_modified_by="Editor",
            modified=datetime(2025, 1, 1, 12, 0, 0),
        )

        document = MagicMock()
        document.paragraphs = _build_paragraphs(["Sample paragraph"])
        document.tables = []
        document.sections = []
        document.core_properties = core

        mock_document.return_value = document

        result = self.processor.process_docx_with_metadata(str(doc_path))

        assert result["text"] == "Sample paragraph"
        assert result["metadata"]["title"] == "Report"
        assert result["metadata"]["author"] == "Editor"
        assert result["metadata"]["modified"] == "2025-01-01T12:00:00"

    @patch("zapomni_core.processors.docx_processor.Document")
    def test_headers_and_footers_included_when_requested(
        self, mock_document: MagicMock, tmp_path: Path
    ) -> None:
        """Should append header/footer content when configured."""
        doc_path = _create_docx_path(tmp_path)

        header = MagicMock()
        header.paragraphs = _build_paragraphs(["Header Info"])

        footer = MagicMock()
        footer.paragraphs = _build_paragraphs(["Footer Info"])

        section = MagicMock()
        section.header = header
        section.footer = footer

        document = MagicMock()
        document.paragraphs = _build_paragraphs(["Body text"])
        document.tables = []
        document.sections = [section]
        document.core_properties = _build_core_properties()

        mock_document.return_value = document

        processor = DOCXProcessor(include_headers_and_footers=True)

        output = processor.process_docx(str(doc_path))

        assert "Header Info" in output
        assert "Footer Info" in output

    def test_process_docx_missing_file_raises_validation_error(self) -> None:
        missing_path = Path("/tmp/nonexistent.docx")

        with pytest.raises(ValidationError) as exc_info:
            self.processor.process_docx(str(missing_path))

        assert "exist" in exc_info.value.message.lower()

    def test_process_docx_wrong_extension_raises_validation_error(self, tmp_path: Path) -> None:
        wrong_path = tmp_path / "file.txt"
        wrong_path.write_text("not a docx")

        with pytest.raises(ValidationError) as exc_info:
            self.processor.process_docx(str(wrong_path))

        assert ".docx" in exc_info.value.message.lower()

    @patch("zapomni_core.processors.docx_processor.Document")
    def test_process_docx_corrupted_file_raises_processing_error(
        self, mock_document: MagicMock, tmp_path: Path
    ) -> None:
        doc_path = _create_docx_path(tmp_path)
        mock_document.side_effect = PackageNotFoundError("corrupted")

        with pytest.raises(ProcessingError) as exc_info:
            self.processor.process_docx(str(doc_path))

        assert "corrupt" in exc_info.value.message.lower()

    @patch("zapomni_core.processors.docx_processor.Document")
    def test_process_docx_bad_zip_raises_processing_error(
        self, mock_document: MagicMock, tmp_path: Path
    ) -> None:
        doc_path = _create_docx_path(tmp_path)
        mock_document.side_effect = BadZipFile("bad zip")

        with pytest.raises(ProcessingError):
            self.processor.process_docx(str(doc_path))
