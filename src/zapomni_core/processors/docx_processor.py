"""DOCXProcessor - Extracts text and metadata from Word documents."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from zipfile import BadZipFile

from zapomni_core.exceptions import ProcessingError, ValidationError
from zapomni_core.utils import get_logger


Metadata = Dict[str, Any]


class DOCXProcessor:
    """Extracts clean text and metadata from DOCX files."""

    def __init__(self, include_headers_and_footers: bool = False) -> None:
        self.logger = get_logger(__name__)
        self.include_headers_and_footers = include_headers_and_footers

    def process_docx(self, file_path: str) -> str:
        """Extract cleaned text from the DOCX document."""
        text, _ = self._process(file_path)
        return text

    def process_docx_with_metadata(self, file_path: str) -> Dict[str, Any]:
        """Return cleaned text together with extracted metadata."""
        text, metadata = self._process(file_path)
        return {"text": text, "metadata": metadata}

    def _process(self, file_path: str) -> Tuple[str, Metadata]:
        path = Path(file_path)
        self.logger.debug("docx_validation_started", file=str(path))
        self._validate_path(path)

        try:
            document = Document(str(path))
            self.logger.info(
                "docx_loaded",
                file=str(path),
                paragraphs=len(getattr(document, "paragraphs", [])),
            )

            paragraph_segments = self._extract_paragraphs(document.paragraphs)
            table_segments = self._extract_tables(document.tables)
            header_footer_segments = []
            if self.include_headers_and_footers:
                header_footer_segments = self._extract_headers_and_footers(document.sections)

            segments = header_footer_segments + paragraph_segments + table_segments
            extracted_text = self._clean_text("\n".join(segments))
            metadata = self._extract_metadata(document, len(paragraph_segments))

            return extracted_text, metadata
        except (PackageNotFoundError, BadZipFile, OSError) as exc:
            self.logger.error(
                "docx_processing_failed",
                file=str(path),
                error=str(exc),
            )
            raise ProcessingError(
                message="Corrupted DOCX file",
                details={"file_path": str(path)},
                original_exception=exc,
            )
        except ProcessingError:
            raise
        except Exception as exc:
            self.logger.error(
                "docx_processing_failed",
                file=str(path),
                error=str(exc),
            )
            raise ProcessingError(
                message=f"Failed to read DOCX: {exc}",
                details={"file_path": str(path)},
                original_exception=exc,
            )

    def _validate_path(self, path: Path) -> None:
        if not path.exists():
            raise ValidationError(
                message=f"DOCX file does not exist: {path}",
                error_code="VAL_004",
                details={"file_path": str(path)},
            )

        if path.suffix.lower() != ".docx":
            raise ValidationError(
                message="File must have .docx extension",
                error_code="VAL_003",
                details={"file_path": str(path)},
            )

    def _extract_paragraphs(self, paragraphs: Iterable[Any]) -> List[str]:
        extracted: List[str] = []
        for paragraph in paragraphs or []:
            raw_text = getattr(paragraph, "text", "")
            if not raw_text:
                continue
            normalized = self._normalize_segment(raw_text)
            if normalized:
                extracted.append(normalized)
        return extracted

    def _extract_tables(self, tables: Iterable[Any]) -> List[str]:
        extracted: List[str] = []
        for table in tables or []:
            for row in getattr(table, "rows", []):
                cells = [getattr(cell, "text", "") for cell in getattr(row, "cells", [])]
                normalized_cells = [self._normalize_segment(cell) for cell in cells if cell.strip()]
                if normalized_cells:
                    extracted.append(" | ".join(normalized_cells))
        return extracted

    def _extract_headers_and_footers(self, sections: Iterable[Any]) -> List[str]:
        extracted: List[str] = []
        for section in sections or []:
            header = getattr(section, "header", None)
            footer = getattr(section, "footer", None)
            extracted.extend(self._extract_header_footer_paragraphs(header))
            extracted.extend(self._extract_header_footer_paragraphs(footer))
        return extracted

    def _extract_header_footer_paragraphs(self, element: Any) -> List[str]:
        if not element:
            return []
        return self._extract_paragraphs(getattr(element, "paragraphs", []))

    def _clean_text(self, text: str) -> str:
        cleaned_lines: List[str] = []
        for line in text.splitlines():
            normalized = self._normalize_segment(line)
            if normalized:
                cleaned_lines.append(normalized)
        return "\n".join(cleaned_lines).strip()

    @staticmethod
    def _normalize_segment(segment: str) -> str:
        return " ".join(segment.split()) if segment else ""

    def _extract_metadata(self, document: Any, paragraph_count: int) -> Metadata:
        core = getattr(document, "core_properties", None)
        if not core:
            return {
                "title": None,
                "author": None,
                "modified": None,
                "paragraph_count": paragraph_count,
            }

        modified = getattr(core, "modified", None)
        modified_value = modified.isoformat() if modified else None
        author = core.author or getattr(core, "last_modified_by", None)

        return {
            "title": getattr(core, "title", None),
            "author": author,
            "modified": modified_value,
            "paragraph_count": paragraph_count,
        }
